import asyncio
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
import aiofiles
from functools import partial

from data_management.document_processor import (
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_txt,
    create_chunks_by_sentence
)
from config import settings

logger = logging.getLogger(__name__)

class AsyncDocumentProcessor:
    def __init__(self):
        # Thread pool для I/O операций
        self.thread_executor = ThreadPoolExecutor(max_workers=settings.max_workers)
        # Process pool для CPU-intensive операций (OCR)
        self.process_executor = ProcessPoolExecutor(max_workers=2)
        self._semaphore = asyncio.Semaphore(settings.max_workers)
    
    async def process_document(self, file_path: Path) -> List[Dict[str, Any]]:
        """Асинхронно обрабатывает один документ"""
        async with self._semaphore:  # Ограничиваем параллельную обработку
            try:
                logger.info(f"Processing document: {file_path.name}")
                
                # Определяем тип файла
                file_ext = file_path.suffix.lower()
                
                # Извлекаем текст асинхронно
                if file_ext == '.docx':
                    text, metadata = await self._process_docx(file_path)
                elif file_ext == '.pdf':
                    text, metadata = await self._process_pdf(file_path)
                elif file_ext == '.txt':
                    text, metadata = await self._process_txt(file_path)
                else:
                    logger.warning(f"Unsupported file type: {file_ext}")
                    return []
                
                if not text or len(text.strip()) < 50:
                    logger.warning(f"Document {file_path.name} has insufficient text")
                    return []
                
                # Создаем метаданные файла
                file_metadata = {
                    "file_name": file_path.name,
                    "file_path": str(file_path),
                    **metadata
                }
                
                # Создаем чанки асинхронно
                chunks = await self._create_chunks_async(text, file_metadata)
                
                logger.info(f"Document {file_path.name} processed: {len(chunks)} chunks")
                return chunks
                
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                return []
    
    async def process_documents_batch(self, file_paths: List[Path]) -> List[Dict[str, Any]]:
        """Обрабатывает батч документов параллельно"""
        tasks = [self.process_document(fp) for fp in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        all_chunks = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"Failed to process {file_paths[i]}: {result}")
            elif result:
                all_chunks.extend(result)
        
        return all_chunks
    
    async def _process_docx(self, file_path: Path) -> tuple[str, dict]:
        """Асинхронная обработка DOCX"""
        loop = asyncio.get_event_loop()
        
        # Читаем файл асинхронно
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()
        
        # Обрабатываем в thread pool
        text = await loop.run_in_executor(
            self.thread_executor,
            partial(self._extract_docx_sync, content)
        )
        
        metadata = {"file_type": "docx"}
        return text, metadata
    
    def _extract_docx_sync(self, content: bytes) -> str:
        """Синхронная обработка DOCX из байтов"""
        import io
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    async def _process_pdf(self, file_path: Path) -> tuple[str, dict]:
        """Асинхронная обработка PDF"""
        loop = asyncio.get_event_loop()
        
        # Обрабатываем в thread pool
        result = await loop.run_in_executor(
            self.thread_executor,
            extract_text_from_pdf,
            str(file_path),
            50  # min_words_per_page
        )
        
        text, metadata = result
        
        # Если нужен OCR, используем process pool
        if metadata.get('ocr_used', False):
            logger.info(f"Using OCR for {file_path.name}")
            text = await loop.run_in_executor(
                self.process_executor,
                self._ocr_pdf,
                str(file_path)
            )
        
        return text, metadata
    
    def _ocr_pdf(self, file_path: str) -> str:
        """OCR обработка в отдельном процессе"""
        # Импорты внутри функции для process pool
        from data_management.document_processor import extract_text_from_pdf_with_ocr
        return extract_text_from_pdf_with_ocr(file_path)
    
    async def _process_txt(self, file_path: Path) -> tuple[str, dict]:
        """Асинхронная обработка TXT"""
        import chardet
        
        # Определяем кодировку
        async with aiofiles.open(file_path, 'rb') as f:
            raw_data = await f.read(10000)  # Читаем первые 10KB
        
        detected = chardet.detect(raw_data)
        encoding = detected['encoding'] or 'utf-8'
        
        # Читаем весь файл с правильной кодировкой
        async with aiofiles.open(file_path, 'r', encoding=encoding, errors='replace') as f:
            text = await f.read()
        
        metadata = {"file_type": "txt", "encoding": encoding}
        return text, metadata
    
    async def _create_chunks_async(self, text: str, file_metadata: dict) -> List[Dict[str, Any]]:
        """Асинхронное создание чанков"""
        loop = asyncio.get_event_loop()
        
        # Выполняем в thread pool
        chunks = await loop.run_in_executor(
            self.thread_executor,
            create_chunks_by_sentence,
            text,
            file_metadata,
            settings.chunk_size,
            settings.min_chunk_size,
            settings.chunk_overlap
        )
        
        return chunks
    
    def __del__(self):
        """Закрываем executors при удалении"""
        self.thread_executor.shutdown(wait=False)
        self.process_executor.shutdown(wait=False)