from fastapi import APIRouter, Request, HTTPException
from fastapi.responses import FileResponse, JSONResponse
import os
import uuid
from pathlib import Path
import aiofiles
from typing import Dict, Any
import logging
from docx import Document
from docxtpl import DocxTemplate

from config import settings

logger = logging.getLogger(__name__)
router = APIRouter()

# Директории для генерации
GENERATED_DIR = Path("tmp/generated")
TEMPLATES_DIR = Path("templates")

# Создаем директории при импорте
GENERATED_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATES_DIR.mkdir(exist_ok=True)

@router.post("/generate")
async def generate_file(request: Request):
    """Генерация документа через LLM"""
    try:
        data = await request.json()
        description = data.get("description", "")
        
        if not description:
            raise HTTPException(status_code=400, detail="Description is required")
        
        # Используем cache для похожих генераций
        cache_manager = request.state.cache
        cache_key = f"generate:{description[:100]}"
        
        if cache_manager:
            cached = await cache_manager.get(cache_key)
            if cached and isinstance(cached, str):
                # Создаем новый файл с тем же содержимым
                filename = f"{uuid.uuid4()}.docx"
                filepath = GENERATED_DIR / filename
                
                doc = Document()
                doc.add_paragraph(cached)
                doc.save(str(filepath))
                
                return {"download_url": f"/api/download/{filename}"}
        
        # Генерация через OpenAI
        from langchain_openai import ChatOpenAI
        
        llm = ChatOpenAI(
            openai_api_key=settings.openai_api_key,
            temperature=0.7,
            model_name=settings.openai_model
        )
        
        prompt = f"""Сгенерируй подробный отчет по следующему описанию задачи. 
        Ответ дай в виде связного текста на русском языке.
        Структурируй ответ с заголовками и подразделами.
        
        Описание: {description}"""
        
        response = await llm.ainvoke(prompt)
        content = response.content.strip()
        
        # Кешируем результат
        if cache_manager:
            await cache_manager.set(cache_key, content, ttl=3600)
        
        # Создаем DOCX файл
        filename = f"{uuid.uuid4()}.docx"
        filepath = GENERATED_DIR / filename
        
        doc = Document()
        
        # Парсим и форматируем контент
        lines = content.split('\n')
        for line in lines:
            if line.startswith('#'):
                # Заголовок
                level = line.count('#')
                text = line.lstrip('#').strip()
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                else:
                    doc.add_heading(text, level=3)
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(str(filepath))
        
        logger.info(f"Generated document: {filename}")
        return {"download_url": f"/api/download/{filename}"}
        
    except Exception as e:
        logger.error(f"Generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/generate_practice_plan")
async def generate_practice_plan(request: Request):
    """Генерация плана практики по шаблону"""
    try:
        data = await request.json()
        
        # Проверяем обязательные поля
        required_fields = ["student_name", "practice_type", "start_date", "end_date", "tasks"]
        for field in required_fields:
            if field not in data:
                raise HTTPException(status_code=400, detail=f"Field '{field}' is required")
        
        # Путь к шаблону
        template_path = TEMPLATES_DIR / "practice_plan_template.docx"
        
        if not template_path.exists():
            # Создаем простой шаблон если его нет
            await create_practice_template(template_path)
        
        # Генерируем документ
        doc = DocxTemplate(str(template_path))
        doc.render(data)
        
        filename = f"practice_plan_{uuid.uuid4()}.docx"
        filepath = GENERATED_DIR / filename
        doc.save(str(filepath))
        
        logger.info(f"Generated practice plan: {filename}")
        return {"download_url": f"/api/download/{filename}"}
        
    except Exception as e:
        logger.error(f"Practice plan generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/generate_practice_plan_text")
async def generate_practice_plan_text(request: Request):
    """Генерация текста плана практики"""
    try:
        data = await request.json()
        
        template = """Календарный план производственной практики

Студент: {student_name}
Вид практики: {practice_type}
Сроки: {start_date} — {end_date}

Задачи:
{tasks}

Руководитель практики: _____________________
Дата составления: _____________________"""
        
        text = template.format(
            student_name=data.get("student_name", ""),
            practice_type=data.get("practice_type", ""),
            start_date=data.get("start_date", ""),
            end_date=data.get("end_date", ""),
            tasks=data.get("tasks", "")
        )
        
        return {"text": text}
        
    except Exception as e:
        logger.error(f"Text generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/download/{filename}")
async def download_file(filename: str):
    """Скачивание сгенерированного файла"""
    filepath = GENERATED_DIR / filename
    
    if not filepath.exists():
        logger.error(f"File not found for download: {filepath}")
        return JSONResponse({"error": "File not found"}, status_code=404)
    
    # Проверяем что файл не слишком старый (защита от накопления файлов)
    if filepath.stat().st_mtime < (Path.ctime(filepath) - 86400):  # 24 часа
        filepath.unlink()
        return JSONResponse({"error": "File expired"}, status_code=404)
    
    logger.info(f"Downloading file: {filepath}")
    return FileResponse(
        str(filepath),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

async def create_practice_template(template_path: Path):
    """Создает базовый шаблон практики"""
    doc = Document()
    
    doc.add_heading('Календарный план производственной практики', 0)
    doc.add_paragraph('')
    doc.add_paragraph('Студент: {{ student_name }}')
    doc.add_paragraph('Вид практики: {{ practice_type }}')
    doc.add_paragraph('Сроки: {{ start_date }} — {{ end_date }}')
    doc.add_paragraph('')
    doc.add_heading('Задачи:', level=1)
    doc.add_paragraph('{{ tasks }}')
    
    doc.save(str(template_path))
    logger.info(f"Created practice template: {template_path}")

# Периодическая очистка старых файлов
async def cleanup_old_files():
    """Удаляет файлы старше 24 часов"""
    try:
        cutoff_time = Path.ctime(Path()) - 86400  # 24 часа назад
        
        for filepath in GENERATED_DIR.glob("*.docx"):
            if filepath.stat().st_mtime < cutoff_time:
                filepath.unlink()
                logger.info(f"Deleted old file: {filepath.name}")
                
    except Exception as e:
        logger.error(f"Cleanup error: {e}")